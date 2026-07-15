"""
Genera el documento Word IA_EN_EL_PROYECTO.docx con memoria técnica completa
sobre el sistema de RL y toma de decisiones del proyecto.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(paragraph, text, level=1):
    paragraph.text = text
    paragraph.style = f'Heading {level}'

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading_elm)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    return table

# ── PORTADA ───────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('IA en el Proyecto: Memoria Técnica')
run.bold = True
run.font.size = Pt(22)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle_p.add_run(
    'Sistema de Aprendizaje por Refuerzo (RL) y Toma de Decisiones\n'
    'para Dispatch de Conductores y Rebalanceo de Flota'
)
run2.font.size = Pt(13)
run2.italic = True

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.add_run('Proyecto: City2Cruise — Transport & Lockers Barcelona\n'
               'Versión: 1.0  |  Fecha: mayo 2026\n'
               'Autor técnico: Luis Guillén Servera')

doc.add_page_break()

# ── ÍNDICE (manual) ───────────────────────────────────────────────────────────
add_h1(doc, 'Índice')
for item in [
    '1. Motivación: por qué hace falta RL aquí',
    '2. Arquitectura general del sistema de IA',
    '3. Pipeline de telemetría: del GPS al StateTensor',
    '4. Entorno Gymnasium: CruiseDispatchEnv',
    '5. Agente PPO: arquitectura, entrenamiento e inferencia',
    '6. Integración con el backend: RLDispatchService',
    '7. Digital Twin: escenarios y sim-to-real',
    '8. Rebalanceo autónomo: ReassignmentService',
    '9. Torre de Control: intervención humana',
    '10. Suite de validación y criterios de promoción',
    '11. Infraestructura: Docker y CI/CD',
    '12. Flujo end-to-end completo',
    '13. Variables de entorno y feature flags',
    '14. KPIs y métricas clave',
    '15. Limitaciones y lo que la IA no hace',
]:
    add_bullet(doc, item)

doc.add_page_break()

# ── 1. MOTIVACIÓN ─────────────────────────────────────────────────────────────
add_h1(doc, '1. Motivación: por qué hace falta RL aquí')

add_body(doc,
    'El problema central del proyecto no es simplemente localizar al conductor más cercano. '
    'La decisión de dispatch óptimo combina de forma simultánea múltiples variables '
    'que interactúan entre sí y cambian en tiempo real:'
)
for bullet in [
    'Posición GPS suavizada con filtro de Kalman (no el punto crudo, con ruido)',
    'ETA real en función de velocidad, tráfico y distancia haversine',
    'Urgencia derivada del horario de los cruceros (minutos al "all-aboard")',
    'Ocupación actual de lockers en la zona de servicio',
    'Densidad de demanda por clusters (DBSCAN sobre solicitudes activas)',
    'Disponibilidad real de cada conductor (asignado, en ruta, libre)',
    'Hora del día y patrón de tráfico urbano',
]:
    add_bullet(doc, bullet)

add_body(doc,
    'Una regla estática o greedy (el conductor más cercano) funciona en condiciones simples, '
    'pero se degrada ante picos de demanda, llegada simultánea de varios cruceros, '
    'ruido GPS en zonas portuarias o cambios de última hora. El RL permite aprender '
    'políticas de decisión que mejoran el baseline greedy sin necesidad de modelar '
    'explícitamente cada caso.'
)

add_h2(doc, '1.1 Casos donde el greedy falla')
add_table(doc,
    ['Situación', 'Greedy', 'RL'],
    [
        ['2 cruceros docking simultáneamente', 'Asigna el conductor más cercano a cada solicitud sin priorizar urgencia', 'Prioriza solicitudes con mayor urgency score (minutos al all-aboard)'],
        ['Conductor rápido lejano vs conductor lento cercano', 'Elige siempre el cercano', 'Evalúa ETA real (distancia / velocidad)'],
        ['Zona saturada con muchas solicitudes', 'Asigna conductores sub-óptimamente al no ver el cluster completo', 'Observa demandClusters y distribuye mejor'],
        ['GPS con outliers por interferencia portuaria', 'Puede asignar a un conductor que en realidad está en posición errónea', 'Usa posición Kalman-suavizada con sigma de incertidumbre'],
    ]
)

doc.add_page_break()

# ── 2. ARQUITECTURA GENERAL ───────────────────────────────────────────────────
add_h1(doc, '2. Arquitectura general del sistema de IA')

add_body(doc,
    'El sistema de IA se distribuye en tres microservicios que cooperan de forma '
    'asíncrona: el backend TypeScript (Express 5), el servicio RL (FastAPI + Python) '
    'y el Digital Twin (FastAPI + Python). El frontend React actúa como capa de '
    'observabilidad y control.'
)

add_h2(doc, '2.1 Mapa de componentes')
add_table(doc,
    ['Componente', 'Tecnología', 'Puerto', 'Responsabilidad'],
    [
        ['backend', 'TypeScript + Express 5', '9000', 'API REST, telemetría, dispatch, eventos socket.io'],
        ['rl_service', 'Python + FastAPI + SB3', '8080', 'Inferencia PPO, entrenamiento, rankings'],
        ['digital_twin', 'Python + FastAPI', '8090', 'Simulación de escenarios, estado sintético'],
        ['cruise-connect (frontend)', 'React + Leaflet', '5173 (dev)', 'Mapa en tiempo real, Torre de Control, rankings'],
        ['PostgreSQL + PostGIS', 'PostgreSQL 15', '5432', 'Persistencia de solicitudes, conductores, telemetría'],
    ]
)

add_h2(doc, '2.2 Flujo de decisión de alto nivel')
for step in [
    '1. El backend recibe eventos GPS de conductores → KalmanFilter suaviza la trayectoria.',
    '2. StateFusion ensambla el StateTensor (< 200 ms): posiciones, ETAs, clusters, urgencia, lockers.',
    '3. RLDispatchService envía el tensor al rl_service vía POST /assign (timeout: 2 s).',
    '4. El agente PPO ejecuta inferencia (< 20 ms): devuelve rankings de conductores ordenados por score.',
    '5. GeoDispatchService usa el ranking RL como primera opción; fallback a distancia geo si falla.',
    '6. rebalanceFleetJob evalúa solicitudes stale y, si RL_REBALANCE_ACTIVE=true, llama a ReassignmentService.',
    '7. TwinSyncService replica los eventos al Digital Twin (fire-and-forget, timeout 500 ms).',
    '8. La Torre de Control muestra rankings en vivo y permite intervención manual si es necesario.',
]:
    add_numbered(doc, step)

doc.add_page_break()

# ── 3. PIPELINE DE TELEMETRÍA ─────────────────────────────────────────────────
add_h1(doc, '3. Pipeline de telemetría: del GPS al StateTensor')

add_body(doc,
    'El pipeline de telemetría (Hito 3.D) está implementado en '
    'backend/src/services/telemetry/ y consta de tres módulos que se ejecutan '
    'en secuencia con un objetivo de latencia total < 200 ms.'
)

add_h2(doc, '3.1 KalmanFilter (KalmanFilter.ts)')
add_body(doc,
    'El filtro de Kalman es el primer paso del pipeline. Recibe los últimos 20 puntos '
    'GPS de cada conductor y los procesa para obtener:'
)
for item in [
    'Posición suavizada (lat, lon) con reducción de ruido GPS',
    'Velocidad estimada (vLat, vLon en deg/s)',
    'Incertidumbre 1-sigma en metros (sigmaM)',
]:
    add_bullet(doc, item)
add_body(doc,
    'El estado del filtro es un vector 4D: [lat, lon, vLat, vLon]. La matriz de '
    'covarianza Q modela el ruido del proceso (aceleración urbana típica) y R el '
    'ruido de medición GPS (típicamente 5-15 m en zona portuaria).'
)

add_h2(doc, '3.2 FeatureEngineering (FeatureEngineering.ts)')
add_body(doc, 'Calcula tres features clave en paralelo:')

add_h3(doc, 'computeDemandDensity — Clusters DBSCAN')
add_body(doc,
    'Aplica DBSCAN sobre las solicitudes activas para identificar zonas de alta '
    'demanda. Cada cluster se representa como (centroidLat, centroidLon, requestCount). '
    'El entorno RL observa hasta 5 clusters (MAX_CLUSTERS=5).'
)

add_h3(doc, 'computeETA — Tiempo de llegada')
add_body(doc,
    'Calcula la ETA de cada conductor a la solicitud objetivo usando distancia '
    'haversine / velocidad Kalman. Normaliza distancia sobre 15 km (MAX_DISTANCE_M) '
    'y produce distanceNorm ∈ [0,1].'
)

add_h3(doc, 'computeUrgency — Urgencia de cruceros')
add_body(doc,
    'Para cada crucero activo calcula minutesToDeadline = (allAboardAt - now) / 60. '
    'El urgency score se mapea a [0,1]: solicitudes con < 30 min al "all-aboard" '
    'reciben urgency > 0.7 y activan el bonus de recompensa en el entorno RL.'
)

add_h2(doc, '3.3 StateFusion (StateFusion.ts)')
add_body(doc, 'Orquesta los pasos anteriores y ensambla el StateTensor final:')
add_code(doc,
'''StateTensor {
  version: "1.0",
  generatedAt: <Unix ms>,
  durationMs: <tiempo pipeline>,
  drivers: DriverState[],          // max 10 conductores
  demandClusters: DemandCluster[], // max 5 clusters
  urgency: UrgencyScore[],         // por crucero activo
  lockers: LockerSummary,          // ocupación total
  activeRequestCount: number       // solicitudes activas
}'''
)
add_body(doc,
    'Todas las coordenadas se normalizan al bounding box del área de servicio '
    '(-15.55, 27.99, -15.35, 28.22 para Las Palmas). El tensor se persiste en '
    'telemetry_state_snapshots (fire-and-forget) para auditoría y reentrenamiento.'
)

doc.add_page_break()

# ── 4. ENTORNO GYMNASIUM ──────────────────────────────────────────────────────
add_h1(doc, '4. Entorno Gymnasium: CruiseDispatchEnv')

add_body(doc,
    'El entorno de RL (rl_service/gym_env.py) implementa la interfaz Gymnasium '
    'estándar y modela el problema de dispatch como un MDP (Markov Decision Process).'
)

add_h2(doc, '4.1 Espacio de observación')
add_body(doc,
    'El espacio de observación es un vector continuo de dimensión 69 (OBS_DIM), '
    'con todos los valores normalizados a [0, 1]:'
)
add_table(doc,
    ['Segmento', 'Dimensión', 'Contenido'],
    [
        ['Conductores (×10)', '50 (10×5)', 'lat_norm, lon_norm, speed_norm, eta_norm, is_available'],
        ['Clusters demanda (×5)', '15 (5×3)', 'lat_norm, lon_norm, count_norm'],
        ['Variables globales', '4', 'locker_occ, max_urgency, active_req_norm, time_of_day'],
        ['TOTAL', '69', 'OBS_DIM = 50 + 15 + 4'],
    ]
)

add_h2(doc, '4.2 Espacio de acción')
add_body(doc,
    'Discrete(10) — el agente selecciona el índice de uno de los conductores '
    'disponibles (0..9). Si el índice supera el número de conductores disponibles, '
    'la acción es inválida y recibe penalización de -10.'
)

add_h2(doc, '4.3 Función de recompensa')
add_body(doc, 'La recompensa por cada asignación válida se calcula como:')
add_code(doc,
'''reward_base = (0.5 × urgency + 0.5 × (1 - eta_norm)) × 100   ∈ [0, 100]
bonus       = +30  si urgency > 0.7 AND eta_norm < 0.25        (urgente atendido rápido)
penalty     = -10  si acción inválida                           (índice fuera de rango)

Reward total ∈ [-10, 130]
  - Máximo 130: solicitud urgente (urgency=1.0) con conductor muy cercano (eta_norm→0)
  - Mínimo -10: acción inválida
  - Baseline greedy: ~50-60 (conductor cercano sin priorizar urgencia)'''
)
add_body(doc,
    'El diseño de la recompensa alinea directamente los objetivos de negocio: '
    'la velocidad de respuesta (eta_norm) y la priorización de pasajeros de crucero '
    'con poco tiempo antes del all-aboard (urgency).'
)

add_h2(doc, '4.4 Ciclo de vida del episodio')
for step in [
    'reset(): genera aleatoriamente N conductores y M solicitudes (N=6-10, M=1-12). Las solicitudes se ordenan por urgencia descendente.',
    'step(a): asigna el conductor available[a] a la solicitud más urgente. El conductor se reposiciona en la ubicación de la solicitud (afecta ETAs futuros).',
    'Terminación: done=True cuando todas las solicitudes están asignadas O truncated=True cuando se supera max_steps=20.',
    'El entorno usa numpy.random.Generator (seed replicable) para reproducibilidad en tests.',
]:
    add_numbered(doc, step)

add_h2(doc, '4.5 TensorEncoder')
add_body(doc,
    'La clase TensorEncoder.encode() convierte el estado del episodio al vector de '
    'observación flat. Los slots más allá del número real de conductores/clusters '
    'se rellenan con ceros (zero-padding). Esta representación fija es necesaria '
    'para las redes neuronales densas de SB3.'
)

doc.add_page_break()

# ── 5. AGENTE PPO ─────────────────────────────────────────────────────────────
add_h1(doc, '5. Agente PPO: arquitectura, entrenamiento e inferencia')

add_body(doc,
    'El agente (rl_service/agent.py) encapsula un modelo PPO de Stable-Baselines3 '
    'con gestión de checkpoints, entrenamiento serializado e inferencia concurrente.'
)

add_h2(doc, '5.1 Hiperparámetros PPO')
add_table(doc,
    ['Parámetro', 'Valor', 'Justificación'],
    [
        ['Política', 'MlpPolicy', 'Observación vectorial (no visual)'],
        ['learning_rate', '1e-4', 'Convergencia estable en espacio de acción pequeño'],
        ['n_steps', '1024', 'Buffer por env antes de actualizar'],
        ['batch_size', '256', 'Balance velocidad/varianza de gradientes'],
        ['n_epochs', '10', 'Pasos de optimización por actualización'],
        ['gamma', '0.99', 'Horizonte largo (episodios de hasta 20 pasos)'],
        ['gae_lambda', '0.95', 'GAE: balance bias-varianza estándar PPO'],
        ['clip_range', '0.2', 'PPO clip estándar'],
        ['ent_coef', '0.005', 'Entropía baja: acción determinista en producción'],
        ['net_arch', '[256, 256]', 'Red 2 capas densas 256 neuronas'],
        ['n_envs (vectorized)', '8', 'Paralelismo en entrenamiento'],
    ]
)

add_h2(doc, '5.2 Persistencia del modelo')
add_body(doc,
    'El modelo se guarda en formato SB3 (.zip) en RL_MODEL_PATH '
    '(default: /tmp/cruise_dispatch_ppo.zip). Junto al zip se escribe un archivo '
    '.meta.json con el modelVersion para verificar compatibilidad en cada carga. '
    'El modelo actual es "ppo-v3-anticipatory".'
)

add_h2(doc, '5.3 Inferencia: get_rankings()')
add_body(doc,
    'El método get_rankings(StateTensorInput) realiza inferencia en < 20 ms:'
)
for step in [
    'Convierte los DriverObservation del tensor a objetos SimDriver internos.',
    'Llama a TensorEncoder.encode() para obtener el vector de observación (OBS_DIM=69).',
    'Extrae probabilidades por acción del policy network PPO usando softmax sobre los logits.',
    'Ordena conductores por probabilidad descendente y devuelve lista AssignmentResult con score y rank.',
]:
    add_numbered(doc, step)

add_body(doc,
    'La inferencia es thread-safe: el policy network de SB3 no tiene estado mutable '
    'durante predict(), por lo que múltiples llamadas concurrentes son seguras. '
    'El entrenamiento sí está serializado mediante threading.Lock.'
)

add_h2(doc, '5.4 Entrenamiento CLI (train.py)')
add_body(doc,
    'Para entrenamientos largos offline se usa el CLI standalone:'
)
add_code(doc, 'python -m rl_service.train --timesteps 500000 --eval-freq 50000')
add_body(doc,
    'Usa EvalCallback de SB3 para guardar el mejor modelo y registrar logs '
    'en /logs (compatible con TensorBoard). El modelo óptimo se guarda en '
    'MODEL_PATH/best/best_model.zip.'
)

doc.add_page_break()

# ── 6. RLDISPATCHSERVICE ──────────────────────────────────────────────────────
add_h1(doc, '6. Integración con el backend: RLDispatchService')

add_body(doc,
    'RLDispatchService (backend/src/services/RLDispatchService.ts) es el cliente '
    'HTTP que conecta el backend TypeScript con el microservicio RL Python. '
    'Es el punto de integración principal y está diseñado para ser fail-safe.'
)

add_h2(doc, '6.1 Contrato de integración')
add_table(doc,
    ['Aspecto', 'Detalle'],
    [
        ['Endpoint', 'POST rl_service:8080/assign'],
        ['Payload', 'StateTensor (serializado como JSON)'],
        ['Respuesta', '{ rankings: [{driverId, score, rank, etaMs}], modelVersion, inferenceMs }'],
        ['Feature flag', 'RL_ROUTING_ENABLED=true (default false en producción)'],
        ['Timeout', 'RL_SERVICE_TIMEOUT_MS=2000 (2 segundos)'],
        ['Fallback', 'Array vacío [] → GeoDispatchService usa distancia geo'],
    ]
)

add_h2(doc, '6.2 Política de fallback')
add_body(doc,
    'La función getRLDriverRanking() NUNCA lanza excepciones. Todos los errores '
    '(timeout, red, respuesta malformada, servicio caído) se capturan y devuelven '
    'array vacío. Esto garantiza que el dispatch siempre continúa aunque el RL '
    'no esté disponible.'
)
add_code(doc,
'''Cascade de dispatch en GeoDispatchService:
  1. Si RL_ROUTING_ENABLED=true → intentar getRLDriverRanking()
     - Si devuelve rankings → usar top-N como candidatos preferentes
     - Si devuelve [] (fallback) → continuar paso 2
  2. Distancia geo (haversine) sobre conductores disponibles
  3. Ofertas emitidas vía socket.io a candidatos seleccionados'''
)

add_h2(doc, '6.3 Uso del ranking en GeoDispatchService')
add_body(doc,
    'El ranking RL es advisory por defecto. GeoDispatchService utiliza los '
    'driverIds del ranking para priorizar a quién enviar la oferta de pickup, '
    'pero el conductor sigue pudiendo aceptar o rechazar. El servicio de '
    'reasignación (ReassignmentService) es el que toma acción activa.'
)

doc.add_page_break()

# ── 7. DIGITAL TWIN ───────────────────────────────────────────────────────────
add_h1(doc, '7. Digital Twin: escenarios y sim-to-real')

add_body(doc,
    'El Digital Twin (digital_twin/) es un microservicio FastAPI que simula el '
    'estado operativo de la flota y genera escenarios sintéticos para entrenar y '
    'validar el agente. Implementa el Hito 5.4.'
)

add_h2(doc, '7.1 Endpoints del Twin')
add_table(doc,
    ['Endpoint', 'Método', 'Función'],
    [
        ['GET /health', 'GET', 'Liveness probe'],
        ['GET /state', 'GET', 'Estado actual simulado (conductores, solicitudes, lockers)'],
        ['GET /state/aggregates', 'GET', 'Métricas agregadas (avg_match_seconds, p95, etc.)'],
        ['POST /scenario/run', 'POST', 'Ejecutar escenario sintético con parámetros dados'],
        ['POST /sync', 'POST', 'Recibir evento del backend (fire-and-forget)'],
    ]
)

add_h2(doc, '7.2 Modelo de tráfico (traffic.py)')
add_body(doc,
    'La función traffic_multiplier(hour, weekday) devuelve un factor > 1 cuando '
    'el tráfico ralentiza los tiempos de asignación. Modela las horas punta urbanas:'
)
add_table(doc,
    ['Período', 'Multiplicador', 'Impacto en ETA'],
    [
        ['7h-9h laborable (mañana)', '1.7×', 'ETA +70% por tráfico'],
        ['17h-19h laborable (tarde)', '1.8×', 'ETA +80% por tráfico'],
        ['22h-5h (noche)', '0.85×', 'ETA -15% (tráfico fluido)'],
        ['Fin de semana 11h-22h', '1.1×', 'Tráfico turístico moderado'],
        ['Resto de horas', '1.0×', 'Baseline'],
    ]
)

add_h2(doc, '7.3 Cruise manifest loader (cruise_schedule.py)')
add_body(doc,
    'Carga fixtures JSON con el horario de cruceros del escenario. Cada entrada '
    'incluye vessel_name, scheduled_arrival (ISO 8601 UTC) y all_aboard. '
    'La función active_at(manifest, t) filtra los cruceros atracados en el '
    'instante t para calcular presión de urgencia.'
)
add_body(doc, 'Escenario implementado (scenarios/las_palmas_baseline.json):')
add_bullet(doc, 'AIDAnova: 07:30–16:30, capacidad 6.300 pasajeros')
add_bullet(doc, 'MSC Bellissima: 09:00–18:00, capacidad 4.500 pasajeros')
add_bullet(doc, 'Mein Schiff 6: 11:30–20:00, capacidad 2.500 pasajeros')
add_bullet(doc, 'Configuración: 20 lockers, 15 conductores, 8 horas de operación')

add_h2(doc, '7.4 Pipeline sim-to-real (TwinClient + train_with_twin_scenarios)')
add_body(doc,
    'El endpoint POST /train_from_twin del rl_service encadena:'
)
for step in [
    'Verificar que el Digital Twin está accesible (GET /health).',
    'Ejecutar N escenarios con parámetros variables (seed incremental para reproducibilidad).',
    'Calcular timesteps proporcionales al volumen simulado: max(2000, total_requests × 100).',
    'Disparar agent.train(timesteps) con los datos acumulados.',
    'Devolver resumen: n_scenarios, total_simulated_requests, train_timesteps, train_metrics.',
]:
    add_numbered(doc, step)

add_h2(doc, '7.5 TwinClient y adapter MiroFish')
add_body(doc,
    'TwinClient (rl_service/twin_bridge.py) es una fachada que selecciona entre '
    'el twin interno y MiroFish según la variable TWIN_PROVIDER:'
)
add_table(doc,
    ['TWIN_PROVIDER', 'Implementación', 'Estado'],
    [
        ['internal (default)', '_InternalTwinClient → HTTP al digital_twin propio', 'Productivo'],
        ['mirofish', 'MiroFishTwinAdapter → HTTP a MiroFish API con auth', 'Integración futura'],
    ]
)
add_body(doc,
    'La interfaz es idéntica en ambos casos: health(), get_state(), get_aggregates(), '
    'run_scenario(). Solo cambia la URL base y el esquema de autenticación.'
)

add_h2(doc, '7.6 TwinSyncService (backend)')
add_body(doc,
    'TwinSyncService (backend/src/services/twin/TwinSyncService.ts) envía eventos '
    'del backend al twin de forma fire-and-forget con timeout de 500 ms:'
)
add_table(doc,
    ['Evento', 'Cuándo se emite'],
    [
        ['driver.position_changed', 'Cada vez que el backend acepta una posición GPS válida'],
        ['request.created', 'Al crear una nueva solicitud de pickup'],
        ['request.assigned', 'Al asignar un conductor a una solicitud'],
        ['request.completed', 'Al completar una entrega'],
    ]
)
add_body(doc,
    'Si TWIN_SYNC_ENABLED=false o el twin no responde, el sync falla silenciosamente '
    'sin afectar al flujo de dispatch del backend.'
)

doc.add_page_break()

# ── 8. REBALANCEO AUTÓNOMO ────────────────────────────────────────────────────
add_h1(doc, '8. Rebalanceo autónomo: ReassignmentService')

add_body(doc,
    'El rebalanceo activo (Hito 3.5) transforma el sistema de advisory a actuator: '
    'cuando el RL detecta que existe una asignación mejor, puede cancelar las '
    'ofertas vigentes y redirigir la solicitud a los conductores óptimos.'
)

add_h2(doc, '8.1 ReassignmentService (backend/src/services/ReassignmentService.ts)')
add_body(doc, 'La función reassignRequest({ requestId, newCandidateIds }) realiza:')
for step in [
    'SELECT FOR UPDATE en pickup_requests para obtener lock transaccional.',
    'Verificar que la solicitud está en estado REQUESTED (si ya está ASSIGNED, abortar con not_in_requested_state).',
    'Obtener ofertas PENDING activas (pickup_offers WHERE request_id AND status=PENDING).',
    'Cancelar las ofertas antiguas (UPDATE SET status=CANCELLED_BY_REBALANCE).',
    'Emitir pickup:offer:cancelled vía socket.io a los conductores cancelados.',
    'Emitir new:pickup:request vía socket.io a los newCandidateIds (top-N del RL).',
    'Devolver { reassigned: true, cancelledOfferCount, newCandidateCount }.',
]:
    add_numbered(doc, step)

add_body(doc,
    'El bloqueo FOR UPDATE previene condiciones de carrera: si otro proceso '
    'asigna la solicitud entre el check y la cancelación, la transacción aborta.'
)

add_h2(doc, '8.2 rebalanceFleetJob — gate por feature flag')
add_body(doc,
    'El job de rebalanceo (backend/src/jobs/rebalanceFleetJob.ts) se ejecuta '
    'periódicamente y evalúa solicitudes stale (sin aceptación en STALE_THRESHOLD_MS).'
)
add_table(doc,
    ['Flag', 'Comportamiento'],
    [
        ['RL_REBALANCE_ACTIVE=false (default)', 'Solo emite dispatch:rebalance:suggested (advisory, sin acción)'],
        ['RL_REBALANCE_ACTIVE=true', 'Además llama a reassignRequest() para las solicitudes stale con top-3 del RL'],
    ]
)
add_body(doc,
    'Este gate permite activar el rebalanceo activo de forma gradual en staging '
    'antes de habilitarlo en producción.'
)

doc.add_page_break()

# ── 9. TORRE DE CONTROL ───────────────────────────────────────────────────────
add_h1(doc, '9. Torre de Control: intervención humana')

add_body(doc,
    'La página ControlTowerPage.tsx del frontend proporciona visibilidad completa '
    'del estado del sistema y permite override manual cuando el operador detecta '
    'situaciones que la IA no puede manejar correctamente.'
)

add_h2(doc, '9.1 Componentes de la Torre')
add_table(doc,
    ['Componente', 'Función'],
    [
        ['Mapa Leaflet (tiempo real)', 'Posiciones de conductores (verde=libre, rojo=ocupado), solicitudes activas, lockers'],
        ['RLRankingTable', 'Tabla en vivo de rankings RL por solicitud activa (se actualiza via socket.io)'],
        ['ManualInterventionPanel', 'Botones de acción para el operador (ver 9.2)'],
        ['KPIs dashboard', 'p95_match_seconds, mean_reward, inference_ms_p95 en tiempo real'],
    ]
)

add_h2(doc, '9.2 Acciones de intervención manual')
add_body(doc,
    'Las acciones pasan por el endpoint POST /api/admin/intervention/* con '
    'requireAdmin middleware (JWT con rol admin):'
)
add_table(doc,
    ['Acción', 'Endpoint', 'Qué hace'],
    [
        ['Cancelar request', 'POST /api/admin/intervention/cancel', 'Cancela la solicitud y emite evento al twin'],
        ['Forzar asignación', 'POST /api/admin/intervention/force-assign', 'Asigna un conductor específico ignorando el RL'],
        ['Rebalanceo manual', 'POST /api/admin/intervention/rebalance', 'Llama a reassignRequest() con candidatos elegidos manualmente'],
    ]
)
add_body(doc,
    'Todas las intervenciones quedan registradas en el audit log del backend '
    'con timestamp, usuario admin y tipo de acción.'
)

doc.add_page_break()

# ── 10. SUITE DE VALIDACIÓN ───────────────────────────────────────────────────
add_h1(doc, '10. Suite de validación y criterios de promoción')

add_body(doc,
    'La suite de validación (Hito 6.5) implementada en rl_service/validation/ '
    'compone tres evaluadores que conforman el release gate de IA.'
)

add_h2(doc, '10.1 Evaluador de convergencia (convergence.py)')
add_body(doc,
    'Lee la serie temporal de rewards desde logs TensorBoard o CSV y calcula:'
)
add_table(doc,
    ['Métrica', 'Cálculo', 'Criterio de aprobación'],
    [
        ['mean_reward', 'Media del último window (default 100 episodios)', '> media greedy baseline'],
        ['reward_std', 'Desviación estándar del tail window', 'Referencial'],
        ['coeff_var', 'reward_std / |mean_reward|', '< 0.10 (10% CV)'],
        ['improvement_pct', '(tail_mean - head_mean) / |head_mean|', '≥ 0.0 (no regresión)'],
        ['is_converged', 'coeff_var < 0.10 AND improvement_pct ≥ 0', 'True para promover'],
    ]
)

add_h2(doc, '10.2 Evaluador de fidelidad sim-vs-real (fidelity.py)')
add_body(doc,
    'Compara las métricas del Digital Twin con las de producción para medir '
    'el "reality gap". Si el twin es muy diferente a producción, el entrenamiento '
    'en simulación no generaliza bien.'
)
add_table(doc,
    ['Métrica', 'Fórmula', 'Umbral (pass)'],
    [
        ['delta_avg_pct', '|prod_avg - twin_avg| / twin_avg', '< 20%'],
        ['delta_p95_pct', '|prod_p95 - twin_p95| / twin_p95', '< 20%'],
        ['pass', 'delta_avg_pct < 0.20 AND delta_p95_pct < 0.20', 'True → twin es representativo'],
    ]
)

add_h2(doc, '10.3 Evaluador de robustez (robustness.py)')
add_body(doc,
    'Inyecta perturbaciones en el pipeline (packet loss + ruido GPS) y mide '
    'cuántas posiciones el pipeline logra recuperar/interpolar:'
)
add_bullet(doc, 'inject_packet_loss(points, loss_rate=0.10): elimina aleatoriamente el 10% de paquetes GPS')
add_bullet(doc, 'inject_gps_noise (de synthetic_data.py): añade ruido Gaussiano (sigma=15m) más outliers')
add_bullet(doc, 'evaluate_robustness(points): mide recovered_pct ≥ 0.90 (90% de posiciones recuperadas)')

add_h2(doc, '10.4 Script release gate (scripts/validate_ai_release.py)')
add_body(doc,
    'Compone los tres evaluadores y produce un reporte JSON único. El script '
    'finaliza con exit code 0 si todos pasan, exit code 1 si alguno falla. '
    'Está integrado en el CI (ai-rl-ci.yml) como gate de release obligatorio.'
)
add_code(doc, 'python scripts/validate_ai_release.py\n# → /tmp/ai_release_report.json')

add_h2(doc, '10.5 Regla de promoción (resumen)')
add_body(doc, 'Una policy solo se despliega a producción si cumple las tres condiciones:')
add_numbered(doc, 'Convergencia: is_converged=True (coeff_var < 0.10, sin regresión vs inicio)')
add_numbered(doc, 'Fidelidad: reality gap < 20% en avg y p95 (el twin representa bien producción)')
add_numbered(doc, 'Robustez: pipeline recovers ≥ 90% bajo 10% packet loss + ruido GPS 15m')

doc.add_page_break()

# ── 11. INFRAESTRUCTURA ───────────────────────────────────────────────────────
add_h1(doc, '11. Infraestructura: Docker y CI/CD')

add_h2(doc, '11.1 Docker Compose')
add_body(doc, 'El stack completo se levanta con docker-compose.yml:')
add_table(doc,
    ['Servicio', 'Imagen', 'Puerto', 'Variables clave'],
    [
        ['backend', 'node:20-alpine', '9000', 'RL_ROUTING_ENABLED, RL_SERVICE_URL, TWIN_SYNC_ENABLED, TWIN_URL'],
        ['rl_service', 'python:3.11-slim (Dockerfile propio)', '8080', 'RL_MODEL_PATH, BACKEND_URL, TWIN_URL, TWIN_PROVIDER'],
        ['digital_twin', 'python:3.11-slim', '8090', 'TWIN_ENV=simulation'],
        ['postgres', 'postgis/postgis:15-3.3', '5432', 'POSTGRES_DB, POSTGRES_USER, POSTGRES_PASSWORD'],
    ]
)
add_code(doc, 'docker compose -f docker-compose.dev.yml up -d\ncurl http://localhost:8080/health  # rl_service\ncurl http://localhost:8090/health  # digital_twin')

add_h2(doc, '11.2 GitHub Actions CI (ai-rl-ci.yml)')
add_body(doc,
    'CI dedicado que se dispara en PRs que tocan archivos RL/twin:'
)
add_table(doc,
    ['Job', 'Qué hace'],
    [
        ['python-tests', 'pytest rl_service/tests/ + digital_twin/tests/ + rl_service/validation/tests/'],
        ['backend-tests', 'npm test con filtros: reassignment, rebalance-active, twin-sync, rl-latency'],
        ['release-gate (opcional)', 'validate_ai_release.py — solo en merge a main'],
    ]
)

doc.add_page_break()

# ── 12. FLUJO END-TO-END ──────────────────────────────────────────────────────
add_h1(doc, '12. Flujo end-to-end completo')

add_body(doc, 'Flujo completo desde evento GPS hasta asignación de conductor con RL activo:')
add_code(doc,
'''[Driver GPS event]
       │
       ▼
 Backend: GpsValidationService.acceptGps()
   → KalmanFilter.update(lat, lon)
   → TwinSyncService.syncToTwin("driver.position_changed") [fire-and-forget]
       │
       ▼
 Backend: RequestService.createPickupRequest()
   → INSERT pickup_requests
   → TwinSyncService.syncToTwin("request.created") [fire-and-forget]
       │
       ▼
 Backend: GeoDispatchService.findCandidates()
   → StateFusion.buildStateTensor()
       ├── KalmanFilter × driver
       ├── computeDemandDensity (DBSCAN)
       ├── computeETA (haversine / speed)
       └── computeUrgency (cruise schedules)
   → RLDispatchService.getRLDriverRanking(tensor)
       └── POST http://rl_service:8080/assign
             → RLAgent.get_rankings()
                 → TensorEncoder.encode(OBS_DIM=69)
                 → PPO.policy → softmax(logits)
                 → ranked AssignmentResult[]
   → Si [] (fallback): usar haversine ranking
   → Emitir new:pickup:request via socket.io a top-N drivers
       │
       ▼ (si RL_REBALANCE_ACTIVE=true y request stale)
 Backend: rebalanceFleetJob
   → getRLDriverRanking() → top-3 candidates
   → ReassignmentService.reassignRequest()
       ├── FOR UPDATE lock
       ├── Cancel PENDING offers
       ├── Emit pickup:offer:cancelled (socket.io)
       └── Emit new:pickup:request (socket.io)
       │
       ▼
 Frontend: ControlTowerPage
   → RLRankingTable (subscribe socket)
   → ManualInterventionPanel (si override necesario)'''
)

doc.add_page_break()

# ── 13. VARIABLES DE ENTORNO ──────────────────────────────────────────────────
add_h1(doc, '13. Variables de entorno y feature flags')

add_h2(doc, '13.1 Backend')
add_table(doc,
    ['Variable', 'Default', 'Descripción'],
    [
        ['RL_ROUTING_ENABLED', 'false', 'Activa consulta al rl_service en el dispatch'],
        ['RL_SERVICE_URL', 'http://localhost:8080', 'URL del microservicio RL'],
        ['RL_SERVICE_TIMEOUT_MS', '2000', 'Timeout para /assign (ms)'],
        ['RL_REBALANCE_ACTIVE', 'false', 'Activa rebalanceo activo vs advisory'],
        ['TWIN_SYNC_ENABLED', 'false', 'Activa envío de eventos al Digital Twin'],
        ['TWIN_URL', 'http://localhost:8090', 'URL del Digital Twin'],
        ['TWIN_SYNC_TIMEOUT_MS', '500', 'Timeout para /sync del twin (ms)'],
    ]
)

add_h2(doc, '13.2 RL Service')
add_table(doc,
    ['Variable', 'Default', 'Descripción'],
    [
        ['RL_MODEL_PATH', '/tmp/cruise_dispatch_ppo', 'Ruta prefijo del modelo PPO (.zip)'],
        ['BACKEND_URL', 'http://localhost:9000', 'URL del backend para pull de tensors'],
        ['INTERNAL_KEY', '—', 'X-Internal-Key para endpoints /api/internal del backend'],
        ['TWIN_URL', 'http://localhost:8090', 'URL del Digital Twin para sim-to-real'],
        ['TWIN_PROVIDER', 'internal', 'Proveedor de twin: "internal" o "mirofish"'],
        ['MIROFISH_API_KEY', '—', 'API key para MiroFish (si TWIN_PROVIDER=mirofish)'],
        ['MIROFISH_BASE_URL', '—', 'URL base de MiroFish'],
    ]
)

add_h2(doc, '13.3 Digital Twin')
add_table(doc,
    ['Variable', 'Default', 'Descripción'],
    [
        ['TWIN_ENV', 'simulation', 'Modo de operación del twin'],
    ]
)

doc.add_page_break()

# ── 14. KPIs Y MÉTRICAS ───────────────────────────────────────────────────────
add_h1(doc, '14. KPIs y métricas clave')

add_body(doc,
    'Las métricas de éxito del sistema de IA se definen en términos operativos, '
    'no en términos de accuracy genérica. Son las mismas métricas que se usan '
    'en el release gate y en el dashboard de la Torre de Control.'
)

add_table(doc,
    ['KPI', 'Qué mide', 'Objetivo', 'Cómo se calcula'],
    [
        ['p95_match_seconds', 'Tiempo de asignación en el peor 5% de casos', 'Menor que baseline greedy', 'Percentil 95 de (offer_accepted_at - request_created_at)'],
        ['mean_reward', 'Calidad global de la policy RL', 'Mayor que greedy', 'Media del reward acumulado por episodio en eval set'],
        ['mean_urgency_loss', 'Urgencia no atendida a tiempo', 'No empeorar vs greedy', 'Suma de urgency × 1 donde eta > deadline'],
        ['inference_ms_p95', 'Latencia de inferencia del RL', '< 50 ms (budget real-time)', 'P95 de inferenceMs reportado por /assign'],
        ['twin_fidelity_delta', 'Reality gap simulación vs producción', '< 20% en avg y p95', 'evaluate_fidelity(twin_metrics, prod_metrics)'],
        ['robustness_recovered_pct', 'Resiliencia ante pérdida de paquetes GPS', '≥ 90%', 'evaluate_robustness(perturbed_stream)'],
    ]
)

add_h2(doc, '14.1 Regla de promoción (recap)')
add_body(doc,
    'Solo se despliega una policy nueva si MEJORA mean_reward vs greedy, '
    'NO EMPEORA mean_urgency_loss, y CUMPLE el presupuesto de latencia '
    '(inference_ms_p95 < 50 ms). Esta regla es hard gate en el CI.'
)

doc.add_page_break()

# ── 15. LIMITACIONES ──────────────────────────────────────────────────────────
add_h1(doc, '15. Limitaciones y lo que la IA no hace')

add_h2(doc, '15.1 Lo que la IA no hace')
for item in [
    'No sustituye las reglas de negocio ni el fallback heurístico (geo-distancia). Siempre hay fallback.',
    'No decide sola en producción sin superar el release gate (convergencia + fidelidad + robustez).',
    'No "predice demanda" como fin en sí mismo: cualquier predicción solo es útil si mejora el dispatch.',
    'No gestiona la autenticación, pagos ni lógica de negocio: esa responsabilidad permanece en el backend.',
    'No opera sin observabilidad: sin métricas en Torre de Control, el RL no se activa en producción.',
]:
    add_bullet(doc, item)

add_h2(doc, '15.2 Limitaciones actuales')
add_table(doc,
    ['Limitación', 'Mitigación'],
    [
        ['El twin simula posiciones aleatorias (no carga GPS real histórico)', 'TwinSyncService replica eventos de producción; se irá llenando con datos reales'],
        ['El PPO no tiene memoria temporal (no LSTM)', 'El StateTensor incluye time_of_day y velocidades Kalman como proxies de contexto temporal'],
        ['MiroFish no integrado (adapter placeholder)', 'TwinClient abstrae el proveedor; activar con TWIN_PROVIDER=mirofish cuando esté disponible'],
        ['max 10 conductores en el espacio de observación', 'Suficiente para el área de Las Palmas; escalar OBS_DIM si la flota crece'],
        ['Entrenamiento offline (no online learning en producción)', 'POST /train permite reentrenamiento manual; pipeline de datos sintéticos disponible'],
    ]
)

add_h2(doc, '15.3 Roadmap pendiente')
add_bullet(doc, 'Integración MiroFish (TWIN_PROVIDER=mirofish) — cuando MiroFish esté disponible')
add_bullet(doc, 'Activar RL_ROUTING_ENABLED=true y RL_REBALANCE_ACTIVE=true en staging con datos reales')
add_bullet(doc, 'Implementar reentrenamiento automático periódico (cron semanal con datos de producción)')
add_bullet(doc, 'Ampliar RLRankingTable con histórico de decisiones y auditoría de overrides manuales')

# ── FIN ───────────────────────────────────────────────────────────────────────
doc.add_page_break()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end_p.add_run('— Fin del documento —')
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

output_path = 'docs/IA_EN_EL_PROYECTO.docx'
doc.save(output_path)
print(f"Documento guardado en {output_path}")
